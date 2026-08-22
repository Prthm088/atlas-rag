# mypy: disable-error-code="no-untyped-call,var-annotated,arg-type"

import io
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import PurePath

import pymupdf
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from atlas.errors import AppError

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ALLOWED_TYPES: dict[str, set[str]] = {
    PDF_MIME: {".pdf"},
    DOCX_MIME: {".docx"},
    "text/plain": {".txt"},
    "text/markdown": {".md", ".markdown"},
    "text/html": {".html", ".htm"},
}


@dataclass(slots=True)
class SourceBlock:
    text: str
    page: int | None = None
    section_path: list[str] = field(default_factory=list)


@dataclass(slots=True)
class ParsedDocument:
    blocks: list[SourceBlock]
    page_count: int | None
    metadata: dict[str, object] = field(default_factory=dict)

    @property
    def character_count(self) -> int:
        return sum(len(block.text) for block in self.blocks)


def sanitize_filename(filename: str) -> str:
    name = PurePath(filename.replace("\\", "/")).name.strip()
    name = re.sub(r"[\x00-\x1f\x7f]+", "", name)
    name = re.sub(r"\s+", " ", name).strip()
    if not name or name in {".", ".."}:
        raise AppError("invalid_filename", "Choose a file with a valid name.")
    return name[:240]


def validate_declared_type(filename: str, mime_type: str) -> None:
    suffix = PurePath(filename).suffix.lower()
    allowed_extensions = ALLOWED_TYPES.get(mime_type)
    if not allowed_extensions or suffix not in allowed_extensions:
        raise AppError(
            "unsupported_file_type",
            "Supported files are PDF, DOCX, TXT, Markdown, and HTML.",
            status_code=415,
        )


def _clean_text(value: str) -> str:
    value = value.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _parse_pdf(data: bytes, max_pages: int) -> ParsedDocument:
    if not data.startswith(b"%PDF"):
        raise AppError("invalid_file_content", "The file is not a valid PDF.", status_code=422)
    try:
        document = pymupdf.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise AppError("invalid_file_content", "The PDF could not be opened.", status_code=422) from exc
    try:
        if document.page_count > max_pages:
            raise AppError("page_limit_exceeded", f"PDFs are limited to {max_pages} pages.", status_code=413)
        blocks: list[SourceBlock] = []
        for page_index, page in enumerate(document):
            page_blocks = page.get_text("blocks", sort=True)
            for raw in page_blocks:
                text = _clean_text(str(raw[4]))
                if text:
                    blocks.append(SourceBlock(text=text, page=page_index + 1))
        character_count = sum(len(block.text) for block in blocks)
        minimum_text = max(80, document.page_count * 24)
        if character_count < minimum_text:
            raise AppError(
                "ocr_required",
                "This PDF appears to be scanned. OCR is not available on the free deployment.",
                status_code=422,
            )
        return ParsedDocument(
            blocks=blocks,
            page_count=document.page_count,
            metadata={"format": "pdf"},
        )
    finally:
        document.close()


def _parse_docx(data: bytes) -> ParsedDocument:
    if not data.startswith(b"PK"):
        raise AppError("invalid_file_content", "The file is not a valid DOCX document.", status_code=422)
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            if "word/document.xml" not in archive.namelist():
                raise AppError("invalid_file_content", "The file is not a valid DOCX document.", status_code=422)
        document = DocxDocument(io.BytesIO(data))
    except AppError:
        raise
    except Exception as exc:
        raise AppError("invalid_file_content", "The DOCX file could not be opened.", status_code=422) from exc

    blocks: list[SourceBlock] = []
    headings: list[str] = []
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            match = re.search(r"(\d+)", style_name)
            level = int(match.group(1)) if match else 1
            headings = headings[: max(0, level - 1)] + [text]
        blocks.append(SourceBlock(text=text, section_path=headings.copy()))
    for table in document.tables:
        rows = []
        for row in table.rows:
            values = [_clean_text(cell.text) for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            blocks.append(SourceBlock(text="\n".join(rows), section_path=headings.copy()))
    if not blocks:
        raise AppError("empty_document", "The document does not contain readable text.", status_code=422)
    return ParsedDocument(blocks=blocks, page_count=None, metadata={"format": "docx"})


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise AppError("invalid_text_encoding", "Text files must use UTF-8 encoding.", status_code=422) from exc


def _parse_plain(data: bytes, format_name: str) -> ParsedDocument:
    text = _clean_text(_decode_text(data))
    if not text:
        raise AppError("empty_document", "The document does not contain readable text.", status_code=422)
    blocks = [SourceBlock(text=part) for part in re.split(r"\n\s*\n", text) if part.strip()]
    return ParsedDocument(blocks=blocks, page_count=None, metadata={"format": format_name})


def _parse_markdown(data: bytes) -> ParsedDocument:
    text = _decode_text(data)
    headings: list[str] = []
    blocks: list[SourceBlock] = []
    buffer: list[str] = []

    def flush() -> None:
        value = _clean_text("\n".join(buffer))
        if value:
            blocks.append(SourceBlock(text=value, section_path=headings.copy()))
        buffer.clear()

    for line in text.splitlines():
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*#*$", line)
        if heading:
            flush()
            level = len(heading.group(1))
            title = _clean_text(heading.group(2))
            headings = headings[: level - 1] + [title]
            blocks.append(SourceBlock(text=title, section_path=headings.copy()))
        elif not line.strip():
            flush()
        else:
            buffer.append(line)
    flush()
    if not blocks:
        raise AppError("empty_document", "The document does not contain readable text.", status_code=422)
    return ParsedDocument(blocks=blocks, page_count=None, metadata={"format": "markdown"})


def _parse_html(data: bytes) -> ParsedDocument:
    soup = BeautifulSoup(_decode_text(data), "lxml")
    for element in soup(["script", "style", "noscript", "template", "svg"]):
        element.decompose()
    headings: list[str] = []
    blocks: list[SourceBlock] = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "blockquote"]):
        text = _clean_text(element.get_text(" ", strip=True))
        if not text:
            continue
        if element.name and element.name.startswith("h"):
            level = int(element.name[1])
            headings = headings[: level - 1] + [text]
        blocks.append(SourceBlock(text=text, section_path=headings.copy()))
    if not blocks:
        raise AppError("empty_document", "The document does not contain readable text.", status_code=422)
    return ParsedDocument(blocks=blocks, page_count=None, metadata={"format": "html"})


def parse_document(data: bytes, filename: str, mime_type: str, max_pdf_pages: int) -> ParsedDocument:
    validate_declared_type(filename, mime_type)
    if mime_type == PDF_MIME:
        return _parse_pdf(data, max_pdf_pages)
    if mime_type == DOCX_MIME:
        return _parse_docx(data)
    if mime_type == "text/markdown":
        return _parse_markdown(data)
    if mime_type == "text/html":
        return _parse_html(data)
    return _parse_plain(data, "text")
