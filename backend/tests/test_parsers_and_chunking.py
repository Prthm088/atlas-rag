from io import BytesIO

import pytest
from docx import Document

from atlas.errors import AppError
from atlas.services.chunking import chunk_document
from atlas.services.parsers import ParsedDocument, SourceBlock, parse_document, sanitize_filename


def test_sanitize_filename_removes_paths_and_controls() -> None:
    assert sanitize_filename("../private/\x00  report.md") == "report.md"


def test_rejects_mismatched_declared_type() -> None:
    with pytest.raises(AppError, match="Supported files") as caught:
        parse_document(b"plain text", "notes.exe", "text/plain", 10)
    assert caught.value.code == "unsupported_file_type"
    assert caught.value.status_code == 415


def test_markdown_preserves_heading_context() -> None:
    parsed = parse_document(
        b"# Project\n\nOverview text.\n\n## Risks\n\nA concrete risk.",
        "brief.md",
        "text/markdown",
        10,
    )
    assert parsed.metadata == {"format": "markdown"}
    assert parsed.blocks[-1].section_path == ["Project", "Risks"]
    assert parsed.blocks[-1].text == "A concrete risk."


def test_html_discards_executable_content() -> None:
    parsed = parse_document(
        b"<h1>Policy</h1><script>steal()</script><p>Approved content.</p>",
        "policy.html",
        "text/html",
        10,
    )
    combined = " ".join(block.text for block in parsed.blocks)
    assert "Approved content" in combined
    assert "steal" not in combined


def test_docx_parser_handles_headings_and_tables() -> None:
    document = Document()
    document.add_heading("Operations", level=1)
    document.add_paragraph("Escalate incidents within ten minutes.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Severity"
    table.cell(0, 1).text = "Critical"
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_document(
        buffer.getvalue(),
        "runbook.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        10,
    )
    assert any(block.section_path == ["Operations"] for block in parsed.blocks)
    assert any("Severity | Critical" in block.text for block in parsed.blocks)


def test_invalid_utf8_is_rejected() -> None:
    with pytest.raises(AppError) as caught:
        parse_document(b"\xff\xfe\xfa", "notes.txt", "text/plain", 10)
    assert caught.value.code == "invalid_text_encoding"


def test_chunking_is_deterministic_and_retains_locations() -> None:
    parsed = ParsedDocument(
        blocks=[
            SourceBlock("Alpha sentence. " * 20, page=1, section_path=["Alpha"]),
            SourceBlock("Beta sentence. " * 20, page=2, section_path=["Beta"]),
        ],
        page_count=2,
    )
    first = chunk_document(parsed, max_tokens=40, overlap_tokens=5)
    second = chunk_document(parsed, max_tokens=40, overlap_tokens=5)

    assert first
    assert [chunk.content_hash for chunk in first] == [chunk.content_hash for chunk in second]
    assert [chunk.index for chunk in first] == list(range(len(first)))
    assert first[0].page_start == 1
    assert first[-1].page_end == 2
    assert all(chunk.token_count <= 45 for chunk in first)
